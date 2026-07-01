from pathlib import Path
from typing import List

from docx import Document
from docx.shared import Inches


class ThesisReportGenerator:
    def __init__(self, title: str = "Thesis Toolkit Report") -> None:
        self.document = Document()
        self.document.add_heading(title, level=1)

    def add_description(self, text: str) -> None:
        self.document.add_paragraph(text)

    def add_summary_table(self, summary_rows: List[dict]) -> None:
        if not summary_rows:
            return
        columns = list(summary_rows[0].keys())
        table = self.document.add_table(rows=1, cols=len(columns))
        hdr_cells = table.rows[0].cells
        for idx, column in enumerate(columns):
            hdr_cells[idx].text = column

        for row in summary_rows:
            cells = table.add_row().cells
            for idx, column in enumerate(columns):
                cells[idx].text = str(row.get(column, ""))

    def add_image(self, image_path: str, caption: str = "") -> None:
        if Path(image_path).exists():
            self.document.add_paragraph(caption)
            self.document.add_picture(image_path, width=Inches(6))

    def save(self, output_path: str) -> str:
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(output)
        return str(output)
